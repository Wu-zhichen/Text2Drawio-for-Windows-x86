import base64
import io
import unittest

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from app.documents import AttachmentError, DocumentExtractor


def _attachment(filename: str, raw: bytes) -> dict:
    return {
        "filename": filename,
        "mime_type": "application/octet-stream",
        "data_base64": base64.b64encode(raw).decode("ascii"),
    }


class DocumentExtractorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.extractor = DocumentExtractor(max_chars_per_file=10_000, max_total_chars=20_000)

    def test_extracts_pdf_text_with_page_markers(self) -> None:
        stream = io.BytesIO()
        pdf = canvas.Canvas(stream)
        pdf.drawString(72, 760, "Paper title: Reliable Diagram Agents")
        pdf.drawString(72, 730, "Method: structured planning and deterministic layout")
        pdf.save()
        extracted = self.extractor.extract([_attachment("paper.pdf", stream.getvalue())])
        self.assertIn("[Page 1]", extracted.context)
        self.assertIn("Reliable Diagram Agents", extracted.context)

    def test_extracts_xlsx_sheets_and_values(self) -> None:
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Metrics"
        worksheet.append(["Month", "Sessions", "Resolution"])
        worksheet.append(["June", 270000, 0.924])
        stream = io.BytesIO()
        workbook.save(stream)
        extracted = self.extractor.extract([_attachment("metrics.xlsx", stream.getvalue())])
        self.assertIn("[Sheet: Metrics]", extracted.context)
        self.assertIn("270000", extracted.context)
        self.assertIn("0.924", extracted.context)

    def test_extracts_docx_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_heading("Research Summary", level=1)
        document.add_paragraph("The proposed method reduces connector crossings.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Accuracy"
        table.cell(1, 1).text = "92.4%"
        stream = io.BytesIO()
        document.save(stream)
        extracted = self.extractor.extract([_attachment("summary.docx", stream.getvalue())])
        self.assertIn("Research Summary", extracted.context)
        self.assertIn("[Table 1]", extracted.context)
        self.assertIn("92.4%", extracted.context)

    def test_scanned_pdf_reports_ocr_limitation(self) -> None:
        from pypdf import PdfWriter

        stream = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        writer.write(stream)
        with self.assertRaisesRegex(AttachmentError, "OCR is not enabled"):
            self.extractor.extract([_attachment("scan.pdf", stream.getvalue())])

    def test_rejects_unsupported_extension(self) -> None:
        with self.assertRaisesRegex(AttachmentError, "Unsupported"):
            self.extractor.extract([_attachment("archive.zip", b"data")])


if __name__ == "__main__":
    unittest.main()
